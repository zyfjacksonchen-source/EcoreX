from __future__ import annotations

import argparse
import json
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Dict

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from common.office_pdf_runtime import (
    OfficePdfRuntimeError,
    build_document_quality_evidence,
    render_document_preview,
)


def _write_json(path: Path, payload: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _render_metadata() -> list[Dict[str, Any]]:
    return [{"page": 1, "artifactRef": "hmac:document-render-001", "width": 1280, "height": 720}]


def _create_clean_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading("Private document title", 0)
    document.add_heading("Private section", level=1)
    document.add_paragraph("Private customer paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Private header"
    table.cell(0, 1).text = "Sensitive header"
    table.cell(1, 0).text = "Hidden value"
    table.cell(1, 1).text = "Internal value"
    document.save(path)


def _break_docx_table_geometry_and_add_orphan_comment(path: Path) -> None:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", namespace)
    with zipfile.ZipFile(path, "r") as archive:
        entries = {name: archive.read(name) for name in archive.namelist()}
    root = ET.fromstring(entries["word/document.xml"])
    for table in root.findall(f".//{{{namespace}}}tbl"):
        for child in list(table):
            if child.tag == f"{{{namespace}}}tblGrid":
                table.remove(child)
        properties = table.find(f"{{{namespace}}}tblPr")
        if properties is not None:
            for child in list(properties):
                if child.tag == f"{{{namespace}}}tblW":
                    properties.remove(child)
        for cell in table.findall(f".//{{{namespace}}}tc"):
            cell_properties = cell.find(f"{{{namespace}}}tcPr")
            if cell_properties is not None:
                for child in list(cell_properties):
                    if child.tag == f"{{{namespace}}}tcW":
                        cell_properties.remove(child)
    entries["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    entries["word/comments.xml"] = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:comment w:id="0" w:author="Private reviewer" w:date="2026-01-01T00:00:00Z">'
        b"<w:p><w:r><w:t>Private orphan comment</w:t></w:r></w:p>"
        b"</w:comment></w:comments>"
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, data in entries.items():
            archive.writestr(name, data)


def _create_bad_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Private body without heading")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Private cell"
    document.save(path)
    _break_docx_table_geometry_and_add_orphan_comment(path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Smoke R24-07 document quality evidence.")
    parser.add_argument("--output", default="docs/v0.2.4/artifacts/document-quality-smoke.json")
    args = parser.parse_args()

    with tempfile.TemporaryDirectory(prefix="ecorex-docx-quality-") as tmp:
        root = Path(tmp)
        clean_path = root / "clean-private-doc.docx"
        bad_path = root / "bad-private-doc.docx"
        _create_clean_docx(clean_path)
        _create_bad_docx(bad_path)

        render_backend = "provided-redacted-metadata"
        actual_render = None
        actual_render_available = False
        try:
            actual_render = render_document_preview(clean_path, root / "rendered", max_pages=1)
            render_backend = str(actual_render.get("renderBackend") or "libreoffice-pdf")
            render_items = actual_render.get("artifacts") or []
            actual_render_available = True
        except OfficePdfRuntimeError as exc:
            render_items = _render_metadata()
            actual_render = {"status": "unavailable", "errorType": type(exc).__name__}

        clean = build_document_quality_evidence(clean_path, renders=render_items, visual_inspection_passed=True)
        bad = build_document_quality_evidence(bad_path, renders=[], visual_inspection_passed=True)

        expected_clean_no_render_failures = {"render-docx"}
        required_bad_failures = {"design-preset", "render-docx", "table-geometry", "redline-preserve"}
        clean_failed_checks = {item["id"] for item in clean.get("checks", []) if item.get("status") == "fail"}
        bad_failed_checks = {item["id"] for item in bad.get("checks", []) if item.get("status") == "fail"}
        synthetic_render_rejected = (
            not actual_render_available
            and "render-docx" in clean_failed_checks
            and not clean.get("renderedArtifacts")
        )
        clean_render_contract_ok = (
            clean.get("status") == "pass" and not clean_failed_checks
            if actual_render_available
            else clean.get("status") == "fail"
            and synthetic_render_rejected
            and clean_failed_checks == expected_clean_no_render_failures
        )

        serialized = json.dumps({"clean": clean, "bad": bad}, ensure_ascii=False)
        leaks = [
            item
            for item in (
                "Private document title",
                "Private customer paragraph",
                "Private body without heading",
                "Private orphan comment",
                "Hidden value",
                str(root),
                clean_path.name,
                bad_path.name,
            )
            if item and item in serialized
        ]
        payload = {
            "status": (
                "PASS"
                if (
                    clean_render_contract_ok
                    and bad.get("status") == "fail"
                    and required_bad_failures <= bad_failed_checks
                    and not leaks
                )
                else "FAIL"
            ),
            "renderBackend": render_backend,
            "actualRender": actual_render,
            "actualRenderAvailable": actual_render_available,
            "syntheticRenderRejected": synthetic_render_rejected,
            "cleanStatus": clean.get("status"),
            "cleanFailedChecks": sorted(clean_failed_checks),
            "expectedCleanNoRenderFailures": sorted(expected_clean_no_render_failures),
            "badStatus": bad.get("status"),
            "badFailedChecks": sorted(bad_failed_checks),
            "requiredBadFailures": sorted(required_bad_failures),
            "cleanSummary": clean.get("documentAnalysis", {}).get("summary", {}),
            "badSummary": bad.get("documentAnalysis", {}).get("summary", {}),
            "leakCount": len(leaks),
            "redacted": True,
        }
        _write_json(Path(args.output), payload)
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return 0 if payload["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
